import random
from datetime import datetime

from django.contrib import messages
from django.core.exceptions import ObjectDoesNotExist
from django.core.files import File
from django.db import IntegrityError
from django.forms.formsets import formset_factory
from django.http import HttpResponseRedirect
from django.shortcuts import render
from django.views.generic import View
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from openpyxl import load_workbook

from repository.forms import (AssignOrRemoveStaffForm, NewSubjectForm,
                              QuestionBankUploadForm,
                              QuestionPaperCategoryForm,
                              QuestionPaperGenerateForm)
from repository.models import Department, Exam, Question, Subject, User
from shared import current_user, is_user_hod


class NewSubject(View):
    """Let's a new subject to be created"""

    template = 'new_subject.html'
    error = ''
    status = 200

    def get(self, request):
        department_list = Department.objects.all()
        return render(request, self.template,
                      {'department_list': department_list})

    def post(self, request):
        department_list = Department.objects.all()
        try:
            form = NewSubjectForm(request.POST)
            if form.is_valid():
                input_code = form.cleaned_data['code']
                input_name = form.cleaned_data['name']
                input_credit = form.cleaned_data['credit']
                input_course = form.cleaned_data['course']
                input_semester = form.cleaned_data['semester']
                input_department = current_user(request).department.id
                subject = Subject(code=input_code, name=input_name,
                                  credit=input_credit, course=input_course,
                                  semester=input_semester,
                                  department_id=input_department)
                subject.save()
                return HttpResponseRedirect('/subject/' + str(subject.id))
        except IntegrityError:
            self.error = 'Subject Code already exists'
        return render(request, self.template,
                      {
                          'department_list': department_list,
                          'error': self.error},
                      status=self.status)


class ViewSubject(View):
    """Display details of a subject"""

    error = ''
    status = 200

    def get(self, request, subject_id):
        try:
            subject = Subject.objects.get(id=subject_id)
            resource_list = subject.resource_set.all()
            subscription_status = True
            is_hod = False
            has_staff = False
            is_staff = False
            subject_staff_list = subject.staff.all()
            if subject_staff_list:
                has_staff = True
            if 'user' in request.session:
                user = current_user(request)
                if subject not in user.subscribedsubjects.all():
                    subscription_status = False
                if user.status == 'hod' and \
                        user.department == subject.department:
                    is_hod = is_user_hod(request, subject)
                if user in subject.staff.all():
                    is_staff = True
            return render(request, 'subject_resource_list.html',
                          {
                              'subject': subject,
                              'resource_list': resource_list,
                              'subscription_status': subscription_status,
                              'is_hod': is_hod,
                              'is_staff': is_staff,
                              'has_staff': has_staff,
                              'subject_staff_list': subject_staff_list
                          })
        except ObjectDoesNotExist:
            self.error = 'The subject you requested does not exist.'
            self.status = 404
            return render(request, 'error.html',
                          {
                              'error': self.error
                          }, status=self.status)

    def post(self, request, subject_id):
        self.error = 'POST Method not supported.'
        self.status = 405
        return render(request, 'error.html',
                      {
                          'error': self.error
                      }, status=self.status)


class SubscribeUser(View):
    """Let's a user subscribe to a subject"""

    error = ""
    status = 200
    template = "error.html"

    def get(self, request, subject_id):
        try:
            user = current_user(request)
            if user:
                subject = Subject.objects.get(id=subject_id)
                subject.students.add(current_user(request))
                subject.save()
                return HttpResponseRedirect('/subject/' + subject_id)
            else:
                self.error = "You are not logged in."
                self.status = 403
        except ObjectDoesNotExist:
            self.error = "Requested subject not found."
            self.status = 404
        return render(request, self.template,
                      {
                          'error': self.error
                      }, status=self.status)


class UnsubscribeUser(View):
    """Let's a user unsubscribe to a subject"""

    error = ""
    status = 200
    template = "error.html"

    def get(self, request, subject_id):
        try:
            subject = Subject.objects.get(id=subject_id)
            if 'user' in request.session:
                user = current_user(request)
                if user in subject.students.all():
                    subject.students.remove(user)
                    subject.save()
                    return HttpResponseRedirect('/subject/' + subject_id)
                else:
                    self.error = 'You are not subscribed to this subject.'
                    self.status = 403
            else:
                self.error = "You are not logged in."
                self.status = 403
        except ObjectDoesNotExist:
            self.error = 'The subject you requested does not exist.'
        return render(request, self.template,
                      {
                          'error': self.status
                      }, status=self.status)


class AssignStaff(View):
    """Assigns staff to a subject. Available to HOD of the subject."""

    error = ""
    template = "assign_staff.html"
    status = 200

    def get(self, request, subject_id):
        subject = Subject.objects.get(id=subject_id)
        is_hod = is_user_hod(request, subject)
        staff_list = {}
        for department in Department.objects.all():
            staff_list[department.name] = [x for
                                           x in department.user_set.all()
                                           if x.status == 'teacher' or
                                           x.status == 'hod']
        return render(request, self.template,
                      {
                          'is_hod': is_hod,
                          'staff_list': staff_list,
                          'subject': subject
                      })

    def post(self, request, subject_id):
        subject = Subject.objects.get(id=subject_id)
        is_hod = is_user_hod(request, subject)
        staff_list = {}
        for department in Department.objects.all():
            staff_list[department.name] = [x for
                                           x in department.user_set.all()
                                           if x.status == 'teacher' or
                                           x.status == 'hod']
        try:
            if is_hod:
                form = AssignOrRemoveStaffForm(request.POST)
                if form.is_valid():
                    for staff_id in form.cleaned_data['staffselect']:
                        staff = User.objects.get(id=staff_id)
                        subject.staff.add(staff)
                else:
                    self.error = 'Something went wrong.'
                    self.status = 500
                    raise
            else:
                self.error = 'You are not an HOD'
                self.status = 403
        except:
            return render(request, self.template,
                          {
                              'is_hod': is_hod,
                              'staff_list': staff_list,
                              'subject': subject,
                              'error': self.error
                          }, status=self.status)
        return HttpResponseRedirect('/subject/' + subject_id)


class RemoveStaff(View):
    """Removes staff from a subject. Available to HOD of the subject."""

    error = ""
    template = "remove_staff.html"
    status = 200

    def get(self, request, subject_id):
        subject = Subject.objects.get(id=subject_id)
        is_hod = is_user_hod(request, subject)
        staff_list = subject.staff.all()
        return render(request, self.template,
                      {
                          'is_hod': is_hod,
                          'staff_list': staff_list,
                          'subject': subject
                      })

    def post(self, request, subject_id):
        subject = Subject.objects.get(id=subject_id)
        is_hod = is_user_hod(request, subject)
        staff_list = subject.staff.all()
        try:
            if is_hod:
                form = AssignOrRemoveStaffForm(request.POST)
                if form.is_valid():
                    for staff_id in form.cleaned_data['staffselect']:
                        staff = User.objects.get(id=staff_id)
                        subject.staff.remove(staff)
                        subject.save()
                else:
                    self.error = 'Something went wrong.'
                    self.status = 500
                    raise
            else:
                self.error = 'You are not an HOD'
                self.status = 403
        except:
            return render(request, self.template,
                          {
                              'is_hod': is_hod,
                              'staff_list': staff_list,
                              'subject': subject,
                              'error': self.error
                          }, status=self.status)
        return HttpResponseRedirect('/subject/' + subject_id)


class UploadQuestionBank(View):
    """Upload a subject's question bank"""

    def read_excel_file(self, excelfilepath, subject):
        """Read excel file which contains question bank and create question objects
        from it."""
        workbook = load_workbook(filename=excelfilepath)
        for row in workbook.worksheets[0].rows:
            try:
                questiontext = row[1].value
                print "Question Text: ", questiontext
                questionmodule = row[2].value
                print "Question Module: ", questionmodule
                questionpart = row[4].value
                print "Question Part: ", questionpart
                questionco = row[3].value
                print "Question CO: ", questionco
                questionlevel = row[5].value
                print "Question Level: ", questionlevel
                question = Question(text=questiontext,
                                    module=questionmodule,
                                    part=questionpart,
                                    co=questionco,
                                    level=questionlevel
                                    )
                question.subject = subject
                question.save()
                print questiontext, questionmodule, questionpart, \
                    questionco, questionlevel
            except Exception as e:
                print("Error")
                print(e)
                pass

    def get(self, request, subject_id):
        subject = Subject.objects.get(id=subject_id)
        return render(request, 'upload_questionbank.html',
                      {'subject': subject,
                       'user': current_user(request)})

    def post(self, request, subject_id):
        subject = Subject.objects.get(id=subject_id)
        form = QuestionBankUploadForm(request.POST, request.FILES)
        if form.is_valid():
            try:
                qbfile = request.FILES['qbfile']
                with open('/tmp/qb.xlsx', 'wb') as destination:
                    for chunk in qbfile.chunks():
                        destination.write(chunk)
                self.read_excel_file('/tmp/qb.xlsx', subject)
                messages.success(request, "Uploaded succesfully")
                return HttpResponseRedirect('/subject/' + subject_id)
            except:
                return render(request, 'upload_questionbank.html',
                              {'subject': subject,
                               'error': 'Some problem with the file',
                               'user': current_user(request)})
        else:
            return render(request, 'upload_questionbank.html',
                          {'subject': subject,
                           'error': 'Some problem with the file',
                           'user': current_user(request)})


class GenerateQuestionPaper(View):
    """Generate subject's question paper"""

    def get(self, request, subject_id):
        error = ''
        subject = Subject.objects.get(id=subject_id)
        QuestionFormSet = formset_factory(QuestionPaperCategoryForm)
        question_categories_set = QuestionFormSet()
        return render(request, 'generatequestionpaper.html',
                      {'subject': subject,
                       'qpformset': question_categories_set,
                       'error': error,
                       'user': current_user(request)})

    def select_random(self, itemlist, count):
        """Select n items randomly from a list. (Reservoir sampling)"""
        result = []
        N = 0
        for item in itemlist:
            N += 1
            if len(result) < count:
                result.append(item)
            else:
                s = int(random.random() * N)
                if s < count:
                    result[s] = item
        return result

    def make_document(self, subject, questions, exam, marks, time):

        print("Questions")
        print(questions)
        today = datetime.today()
        filename = subject.name.replace(' ', '_') + '_' + \
            str(today.day) + str(today.month) + str(today.year)

        document = Document()

        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_text = 'Adi Shankara Institute of Engineering and Technology'
        run = paragraph.add_run(paragraph_text)
        run.bold = True
        font = run.font
        font.name = 'Times New Roman'
        font.size = Pt(14)

        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_text = exam.name
        paragraph.add_run(paragraph_text)

        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_text = subject.name
        paragraph.add_run(paragraph_text)

        paragraph = document.add_paragraph()
        paragraph_format = paragraph.paragraph_format
        length_of_marks = len(str(marks))
        length_of_time = len(str(time))
        number_of_spaces = 45 - \
            (len('Marks : ') + length_of_marks + len('Time : ') + length_of_time)
        paragraph_text = "Marks : " + \
            str(marks) + " " * (115 - number_of_spaces) + "Time : " + str(time)
        paragraph.add_run(paragraph_text)
        for part in ['Part A', 'Part B', 'Part C']:
            count = 1
            if questions[part]:
                paragraph = document.add_paragraph()
                paragraph_format = paragraph.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph_text = part
                run = paragraph.add_run(paragraph_text)
                run.bold = True
                for question in questions[part]:
                    prefix = str(count) + ". "
                    text = prefix + question.text
                    paragraph = document.add_paragraph()
                    paragraph_format = paragraph.paragraph_format
                    paragraph.add_run(text)
                    count = count + 1
        document.save('/tmp/' + filename + '.docx')
        qpinfile = open('/tmp/' + filename + '.docx')
        qpfile = File(qpinfile)
        exam.questionpaper.save(filename + '.docx', qpfile)
        return '/uploads/' + exam.questionpaper.url

    def create_qp_dataset(self, subject, exam, totalmarks, time, criteria):
        """Populates the dataset needed to generate a question paper. Invokes
        make_document() method"""
        questions = {'Part A': [], 'Part B': [], 'Part C': []}
        status = 0
        for trio in criteria:
            module = trio[0]
            part = trio[1]
            level = trio[2]
            count = int(trio[3])
            questiontotallist = Question.objects.filter(
                module=module, part=part, level=level)
            part = "Part " + part
            selectedquestions = self.select_random(
                questiontotallist, count)
            questions[part] = questions[part] + selectedquestions
        if questions:
            for part in questions:
                for question in questions[part]:
                    exam.question_set.add(question)
                    exam.save()
            status = 1
        path = self.make_document(subject, questions, exam, totalmarks, time)
        return status, path

    def post(self, request, subject_id):
        error = ''
        subject = Subject.objects.get(id=subject_id)
        QuestionFormSet = formset_factory(QuestionPaperCategoryForm)
        print(request.POST)
        QPForm = QuestionPaperGenerateForm(request.POST)
        examname = ''
        totalmarks = ''
        time = ''
        if QPForm.is_valid():
            examname = QPForm.cleaned_data['examname']
            totalmarks = QPForm.cleaned_data['totalmarks']
            time = QPForm.cleaned_data['time']
            exam = Exam(name=examname, totalmarks=totalmarks, time=time,
                        subject_id=subject.id)
            exam.save()
        question_categories_set = QuestionFormSet(request.POST)
        print question_categories_set
        if question_categories_set.is_valid():
            print("\n\n\n\n\n\n Form Data \n\n\n\n\n\n\n\n")
            question_criteria = []
            for form in question_categories_set.forms:
                if form.is_valid():
                    module = form.cleaned_data['module']
                    part = form.cleaned_data['part']
                    level = form.cleaned_data['level']
                    count = form.cleaned_data['count']
                    question_criteria.append((module, part, level, count))
            status, path = self.create_qp_dataset(subject, exam,
                                                  totalmarks, time,
                                                  question_criteria)
            return HttpResponseRedirect(path)
        else:
            error = 'Choose some questions.'
            return render(request, 'generatequestionpaper.html',
                          {'subject': subject,
                           'qpformset': question_categories_set,
                           'error': error,
                           'user': current_user(request)})
